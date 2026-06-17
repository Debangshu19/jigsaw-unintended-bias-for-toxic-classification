#!/usr/bin/env python3
"""Build an editable Word version of the presentation from the same section fragments.
Text -> styled paragraphs/headings, table.data -> Word tables, note/keypoint/math -> shaded
boxes, figures -> embedded PNGs (assets/png). Run make_assets.py first to create the PNGs."""
import re
from pathlib import Path
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SECT = ROOT / "sections"; PNG = ROOT / "assets" / "png"
OUT = ROOT / "Raven-Presentation.docx"

ORDER = ["00-cover.html", "02-overview.html", "03-how-we-measure.html", "04-stage1-rnn.html",
         "05-stage2-distilbert.html", "06-stage3-bigger.html", "07-stage4-ensemble.html",
         "08-final-model.html", "09-ravenx.html", "10-roadmap.html", "99-references.html"]
FIG2PNG = {"fig-system": "diagram-system.png", "fig-stage1-rnn": "diagram-stage1-rnn.png",
           "fig-distilbert": "diagram-distilbert.png", "fig-bias-variance": "diagram-bias-variance.png",
           "fig-ensemble": "diagram-ensemble.png", "fig-ravenx": "diagram-ravenx.png"}
FIGW = {"fig-system": 6.6, "fig-stage1-rnn": 6.3, "fig-distilbert": 4.4,
        "fig-bias-variance": 6.3, "fig-ensemble": 6.7, "fig-ravenx": 4.6}
TAGS = {"t-measured": ("MEASURED", RGBColor(0x0a, 0x7a, 0x52)),
        "t-lit": ("LITERATURE", RGBColor(0x5a, 0x36, 0xa8)),
        "t-proposed": ("PROPOSED", RGBColor(0x9a, 0x5a, 0x06))}
BLUE = RGBColor(0x2f, 0x5f, 0xe0); GREY = RGBColor(0x5a, 0x64, 0x73)


# ---- minimal DOM from the fragment ----
class Node:
    __slots__ = ("tag", "attrs", "kids", "text", "parent")
    def __init__(self, tag, attrs=None, text=None):
        self.tag = tag; self.attrs = attrs or {}; self.kids = []; self.text = text; self.parent = None

VOID = {"br", "img", "hr", "meta"}

class Tree(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True); self.root = Node("root"); self.cur = self.root
    def handle_starttag(self, t, a):
        n = Node(t, dict(a)); n.parent = self.cur; self.cur.kids.append(n)
        if t not in VOID: self.cur = n
    def handle_startendtag(self, t, a):
        n = Node(t, dict(a)); n.parent = self.cur; self.cur.kids.append(n)
    def handle_endtag(self, t):
        if t in VOID: return
        node = self.cur
        while node is not self.root and node.tag != t: node = node.parent
        self.cur = node.parent if node is not self.root else self.root
    def handle_data(self, d):
        self.cur.kids.append(Node("#text", text=d))

def cls(n): return n.attrs.get("class", "")

def gettext(n):
    out = []
    def rec(x):
        if x.tag == "#text": out.append(x.text or "")
        else:
            for k in x.kids: rec(k)
    rec(n)
    return re.sub(r"\s+", " ", "".join(out)).strip()

def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexfill); tcPr.append(sh)


def add_inlines(par, node, bold=False, italic=False, mono=False, sub=False):
    for k in node.kids:
        if k.tag == "#text":
            if not k.text: continue
            txt = re.sub(r"\s+", " ", k.text)
            r = par.add_run(txt); r.bold = bold; r.italic = italic; r.font.subscript = sub
            if mono: r.font.name = "Consolas"; r.font.size = Pt(8.6)
        elif k.tag in ("b", "strong"): add_inlines(par, k, True, italic, mono, sub)
        elif k.tag in ("i", "em"): add_inlines(par, k, bold, True, mono, sub)
        elif k.tag == "code": add_inlines(par, k, bold, italic, True, sub)
        elif k.tag == "sub": add_inlines(par, k, bold, italic, mono, True)
        elif k.tag == "br": par.add_run().add_break()
        elif k.tag == "span" and "tag" in cls(k):
            for c, (lbl, col) in TAGS.items():
                if c in cls(k):
                    r = par.add_run(" [" + lbl + "] "); r.bold = True; r.font.size = Pt(7.5); r.font.color.rgb = col
                    break
        else:
            add_inlines(par, k, bold, italic, mono, sub)


def render_table(doc, n):
    head, body = [], []
    for sec in n.kids:
        if sec.tag == "thead":
            head += [[c for c in tr.kids if c.tag in ("th", "td")] for tr in sec.kids if tr.tag == "tr"]
        elif sec.tag == "tbody":
            body += [[c for c in tr.kids if c.tag in ("th", "td")] for tr in sec.kids if tr.tag == "tr"]
    rows = head + body
    if not rows: return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, r in enumerate(rows):
        is_head = ri < len(head)
        for ci in range(ncol):
            cell = t.cell(ri, ci); cell.paragraphs[0].text = ""
            if ci < len(r): add_inlines(cell.paragraphs[0], r[ci])
            if is_head:
                shade(cell, "1F3A66")
                for run in cell.paragraphs[0].runs:
                    run.bold = True; run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    doc.add_paragraph()


def render_box(doc, n, fill):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    cell = t.cell(0, 0); shade(cell, fill); cell.paragraphs[0].text = ""
    add_inlines(cell.paragraphs[0], n)
    doc.add_paragraph()


def render_math(doc, n):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    cell = t.cell(0, 0); shade(cell, "EEF1F6"); cell.paragraphs[0].text = ""
    add_inlines(cell.paragraphs[0], n, mono=True)
    doc.add_paragraph()


def render_figure(doc, n):
    figid = n.attrs.get("id", ""); png = PNG / FIG2PNG.get(figid, "")
    if png.name and png.exists():
        try:
            doc.add_picture(str(png), width=Inches(FIGW.get(figid, 6.3)))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[diagram: {figid}]")
    cap = next((gettext(k) for k in n.kids if k.tag == "figcaption"), None)
    if cap:
        p = doc.add_paragraph(); r = p.add_run(cap); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def render_block(doc, n):
    c = cls(n)
    if n.tag == "#text":
        if n.text and n.text.strip():
            doc.add_paragraph().add_run(re.sub(r"\s+", " ", n.text).strip())
        return
    if n.tag == "div" and "kicker" in c:
        p = doc.add_paragraph(); r = p.add_run(gettext(n).upper()); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = BLUE
        return
    if n.tag == "h2": doc.add_heading(gettext(n), level=1); return
    if n.tag == "h3": doc.add_heading(gettext(n), level=2); return
    if n.tag == "h4":
        p = doc.add_paragraph(); p.add_run(gettext(n)).bold = True; return
    if n.tag == "p":
        add_inlines(doc.add_paragraph(), n); return
    if n.tag in ("ul", "ol"):
        style = "List Bullet" if n.tag == "ul" else "List Number"
        for li in n.kids:
            if li.tag == "li": add_inlines(doc.add_paragraph(style=style), li)
        return
    if n.tag == "table": render_table(doc, n); return
    if n.tag == "div" and ("note" in c or "keypoint" in c):
        render_box(doc, n, "FFF6E8" if "note" in c else "EAF0FF"); return
    if n.tag == "div" and "math" in c: render_math(doc, n); return
    if n.tag == "div" and "figure" in c: render_figure(doc, n); return
    if n.tag == "div" and "cards" in c:
        for card in n.kids:
            if "card" in cls(card):
                for k in card.kids: render_block(doc, k)
        return
    if n.tag == "div" and "cols" in c:
        for col in n.kids:
            if "col" in cls(col):
                for k in col.kids: render_block(doc, k)
        return
    if n.tag in ("div", "section"):
        for k in n.kids: render_block(doc, k)
        return


def render_cover(doc, sec):
    def find(tag=None, klass=None):
        out = []
        def rec(x):
            if (tag is None or x.tag == tag) and (klass is None or klass in cls(x)): out.append(x)
            for k in x.kids: rec(k)
        rec(sec); return out
    h1 = find(tag="h1"); lede = find(klass="lede"); thesis = find(klass="thesis")
    p = doc.add_paragraph(); p.add_run("RAVEN").bold = True
    p.runs[0].font.size = Pt(11); p.runs[0].font.color.rgb = BLUE
    if h1:
        t = doc.add_paragraph(); r = t.add_run(gettext(h1[0])); r.bold = True; r.font.size = Pt(30)
    if lede:
        l = doc.add_paragraph(); r = l.add_run(gettext(lede[0])); r.font.size = Pt(13); r.font.color.rgb = GREY
    if thesis:
        doc.add_paragraph().add_run(gettext(thesis[0])).italic = True
    stats = find(klass="stat")
    if stats:
        t = doc.add_table(rows=1, cols=len(stats)); t.style = "Table Grid"
        for i, st in enumerate(stats):
            num = next((gettext(x) for x in st.kids if "stat-num" in cls(x)), "")
            lab = next((gettext(x) for x in st.kids if "stat-lab" in cls(x)), "")
            cell = t.cell(0, i); shade(cell, "F2F5FB"); cell.paragraphs[0].text = ""
            rn = cell.paragraphs[0].add_run(num + "\n"); rn.bold = True; rn.font.size = Pt(20)
            cell.add_paragraph().add_run(lab).font.size = Pt(8)
    doc.add_paragraph()
    foot = find(klass="cover-foot")
    if foot: doc.add_paragraph(gettext(foot[0]))
    doc.add_page_break()


def main():
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.8); s.top_margin = s.bottom_margin = Inches(0.75)
    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
    for idx, fn in enumerate(ORDER):
        tree = Tree(); tree.feed((SECT / fn).read_text(encoding="utf-8"))
        sec = next((k for k in tree.root.kids if k.tag == "section"), None)
        if sec is None: continue
        if "cover" in cls(sec):
            render_cover(doc, sec); continue
        if idx > 1: doc.add_page_break()
        for k in sec.kids: render_block(doc, k)
    doc.save(OUT)
    print(f"wrote {OUT.name}")


if __name__ == "__main__":
    main()
